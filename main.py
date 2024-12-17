from seleniumbase import SB
from docx import Document
import time
import random

def get_reviews(sb, num_reviews, selection):
    reviews = []
    
    try:
        # Click the "See all reviews" button
        click = 'span:contains("See all reviews")'
        sb.click(click)
        
        # Wait for the review popup to appear
        time.sleep(random.uniform(6.3, 10.8))  # Random delay
          # Random delay
        if selection == 1:
            pass
        if selection == 2:
            click1 = 'span:contains("Star rating")'
            sb.click(click1)
            time.sleep(random.uniform(1.8, 3.8))
            click2 = 'span:contains("1-star")'
            sb.click(click2)
        if selection == 3:
            click1 = 'span:contains("Star rating")'
            sb.click(click1)
            time.sleep(random.uniform(1.8, 3.8))
            click2 = 'span:contains("2-star")'
            sb.click(click2)
        if selection == 4:
            click1 = 'span:contains("Star rating")'
            sb.click(click1)
            time.sleep(random.uniform(1.8, 3.8))
            click2 = 'span:contains("3-star")'
            sb.click(click2)
        if selection == 5:
            click1 = 'span:contains("Star rating")'
            sb.click(click1)
            time.sleep(random.uniform(1.8, 3.8))
            click2 = 'span:contains("4-star")'
            sb.click(click2)
        if selection == 6:
            click1 = 'span:contains("Star rating")'
            sb.click(click1)
            time.sleep(random.uniform(1.8, 3.8))
            click2 = 'span:contains("5-star")'
            sb.click(click2)
        time.sleep(random.uniform(4.3, 8.8))  # Random delay

        
        for num in range(1, num_reviews): 
            # Extract reviews
            element = f'(//div[@class="RHo1pe"]/div[1])[{num}]'
            # review_text = sb.find_element("css selector", element)
            sb.click(element)
            text = sb.get_text(element)
            print(f'\nScraping review number: {num} \n')
            print(text)
            if text not in reviews:  # Avoid duplicates
                reviews.append(text)

            time.sleep(random.uniform(0.8, 1.8))  # Random delay to appear more human-like
    
    except Exception as e:
        print(f"An error occurred while getting reviews: {e}")
    
    return reviews

def save_to_word(reviews, filename="reviews.docx"):
    doc = Document()
    doc.add_heading('Google Play Store Reviews', 0)
    
    for i, review in enumerate(reviews, 1):
        doc.add_paragraph(f"Review {i}:")
        doc.add_paragraph(review)
        doc.add_paragraph()  # Add a blank line between reviews
    
    doc.save(filename)
    print(f"Reviews saved to: {filename} \nPlease Check the folder.")

def main():
    url = input("\nPlease enter the Google Play Store link: ")
    num_reviews = input('\nHow many reviews you want to scrape?\nPut the number only: ')
    selection = int(input('\nSelect the below number from square braket for your rating of reviews.\n[1] Star rating(all of the reviews)\n[2] 1-star\n[3] 2-star\n[4] 3-star\n[5] 4-star\n[6] 5-star\n: '))
    
    with SB(uc=True) as sb:
        try:
            sb.open(url)
            
            # Wait for the page to load
            time.sleep(random.uniform(8.5, 15.8))  # Random delay
            
            # Scrape reviews
            reviews = get_reviews(sb, int(num_reviews)+1, selection)
            
            # Save reviews to a Word document
            save_to_word(reviews)
        
        except Exception as e:
            print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()